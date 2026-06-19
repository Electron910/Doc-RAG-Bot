import os
import docx
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
DATA_DIR = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "data")

# Ensure DATA_DIR exists
os.makedirs(DATA_DIR, exist_ok=True)

# -----------------
# 1. Text Content
# -----------------

txt_content = """Machine learning (ML) is a field of study in artificial intelligence concerned with the development and study of statistical algorithms that can learn from data and generalize to unseen data, and thus perform tasks without explicit instructions. Recently, artificial neural networks have been able to surpass many previous approaches in performance.

Machine learning approaches are traditionally divided into three broad categories, depending on the nature of the "signal" or "feedback" available to the learning system:

Supervised learning: The computer is presented with example inputs and their desired outputs, given by a "teacher", and the goal is to learn a general rule that maps inputs to outputs.
Unsupervised learning: No labels are given to the learning algorithm, leaving it on its own to find structure in its input. Unsupervised learning can be a goal in itself (discovering hidden patterns in data) or a means towards an end (feature learning).
Reinforcement learning: A computer program interacts with a dynamic environment in which it must perform a certain goal (such as driving a vehicle or playing a game against an opponent). As it navigates its problem space, the program is provided feedback that's analogous to rewards, which it tries to maximize.

Other approaches have been developed which don't fit neatly into this three-fold categorisation, and sometimes more than one is used by the same machine learning system. For example, topic modeling, dimensionality reduction, or meta-learning.

The computational analysis of machine learning algorithms and their performance is a branch of theoretical computer science known as computational learning theory. Because training sets are finite and the future is uncertain, learning theory usually does not yield guarantees of the performance of algorithms. Instead, probabilistic bounds on the performance are quite common. The bias-variance decomposition is one way to quantify generalization error.

For the best performance in the context of generalization, the complexity of the hypothesis should match the complexity of the function underlying the data. If the hypothesis is less complex than the function, then the model has underfit the data. If the complexity of the model is increased in response, then the training error decreases. But if the hypothesis is too complex, then the model is subject to overfitting and generalization will be poorer.

In addition to performance bounds, learning theorists study the time complexity and feasibility of learning. In computational learning theory, a computation is considered feasible if it can be done in polynomial time. There are two kinds of time complexity results: Positive results show that a certain class of functions can be learned in polynomial time. Negative results show that certain classes cannot be learned in polynomial time.

Machine learning and data mining often employ the same methods and overlap significantly, but while machine learning focuses on prediction, based on known properties learned from the training data, data mining focuses on the discovery of (previously) unknown properties in the data. Data mining uses many machine learning methods, but with different goals; on the other hand, machine learning also employs data mining methods as "unsupervised learning" or as a preprocessing step to improve learner accuracy. Much of the confusion between these two research communities comes from the basic assumptions they work with: in machine learning, performance is usually evaluated with respect to the ability to reproduce known knowledge, while in knowledge discovery and data mining (KDD) the key task is the discovery of previously unknown knowledge."""

pdf_content = """Quantum computing is a rapidly-emerging technology that harnesses the laws of quantum mechanics to solve problems too complex for classical computers. 

Today, IBM Quantum makes real quantum hardware -- a tool scientists only began to imagine three decades ago -- available to thousands of developers. Our engineers deliver ever-more-powerful superconducting quantum processors at regular intervals, alongside crucial advances in software and quantum-classical orchestration. This work is driving toward the quantum computing era of utility -- a period during which quantum computers can serve as a scientific tool to explore new classes of problems in chemistry, physics, and materials science, well beyond the reach of brute-force classical simulation.

What is a quantum computer?
Classical computers, which include smartphones and laptops, encode information in binary "bits" that can either be 0s or 1s. In a quantum computer, the basic unit of memory is a quantum bit or qubit. Qubits are made using physical systems, such as the spin of an electron or the orientation of a photon.

These systems can be in many different arrangements all at once, a property known as quantum superposition. Qubits can also be inextricably linked together using a phenomenon called quantum entanglement. The result is that a series of qubits can represent different things simultaneously. 

For instance, eight bits is enough for a classical computer to represent any number between 0 and 255. But eight qubits is enough for a quantum computer to represent every number between 0 and 255 at the same time. A few hundred entangled qubits would be enough to represent more numbers than there are atoms in the universe.

This is where quantum computers get their edge over classical ones. In situations where there are a large number of possible combinations, quantum computers can consider them simultaneously. Examples include trying to find the prime factors of a very large number or the best route between two places. 

However, there may also be plenty of situations where classical computers will still outperform quantum ones. Therefore the computers of the future may be a combination of both these types.

The principles of quantum physics that quantum computing is based upon:
Superposition: This is the ability of a quantum system to be in multiple states simultaneously. The go-to example of superposition is the flip of a coin, which consistently lands as heads or tails—a very binary concept. However, when that coin is in mid-air, it is both heads and tails and until it lands, heads and tails simultaneously. Before measurement, the electron exists in quantum superposition.

Entanglement: This is a strong correlation that exists between quantum particles — so strong, in fact, that two or more quantum particles can be inextricably linked in perfect unison, even if separated by great distances. The particles are so intrinsically connected, they can be said to "dance" in instantaneous, perfect unison, even when placed at opposite ends of the universe. This seemingly impossible connection inspired Einstein to describe entanglement as "spooky action at a distance."

Decoherence: This occurs when the quantum state of a qubit collapses due to interaction with its environment. This is one of the biggest challenges in building practical quantum computers, as qubits are incredibly sensitive to heat, electromagnetic fields, and other external disturbances."""

docx_content = """The Roman Empire was the post-Republican period of ancient Rome. As a polity it included large territorial holdings around the Mediterranean Sea in Europe, North Africa, and Western Asia ruled by emperors.

From the accession of Caesar Augustus to the military anarchy of the third century, it was a principate with Italy as metropole of the provinces and the city of Rome as sole capital (27 BC – 286 AD). Although fragmented briefly during the Crisis of the Third Century, the empire was forcibly reassembled, then divided to be ruled by multiple emperors sharing rule over the Western Roman Empire (based in Milan and later in Ravenna) and the Eastern Roman Empire (later also known as the Byzantine Empire; based in Nicomedia and later in Constantinople). Rome remained the nominal capital of both parts until AD 476, when the imperial insignia were sent to Constantinople following the capture of the Western capital of Ravenna by the Germanic barbarians of Odoacer and the subsequent deposition of Romulus Augustulus.

The adoption of Christianity as the state church of the Roman Empire in AD 380 and the fall of the Western Roman Empire to Germanic kings conventionally marks the end of classical antiquity and the beginning of the Middle Ages. Because of these events, along with the gradual Hellenization of the Eastern Roman Empire, historians distinguish the medieval Roman Empire that remained in the Eastern provinces as the Byzantine Empire.

The predecessors of the Roman Empire, the Roman Republic, which had replaced the Roman monarchy in the 6th century BC, became severely destabilized in a series of civil wars and political conflicts. In the mid-1st century BC Julius Caesar was appointed as perpetual dictator and then assassinated in 44 BC. Civil wars and proscriptions continued, culminating in the victory of Octavian, Caesar's adopted son, over Mark Antony and Cleopatra at the Battle of Actium in 31 BC. The following year Octavian conquered the Ptolemaic Kingdom in Egypt.

Octavian's power was then unassailable and in 27 BC the Roman Senate formally granted him overarching power and the new title Augustus, effectively making him the first Roman emperor. The vast Roman territories were organized in senatorial and imperial provinces, except Italy, which continued to serve as the metropole.

The first two centuries of the empire saw a period of unprecedented stability and prosperity known as the Pax Romana ("Roman Peace"). Rome reached its greatest territorial expanse during the reign of Trajan (AD 98–117). A period of increasing trouble and decline began with the reign of Commodus (180–192).

In the 3rd century, the empire underwent a crisis that threatened its existence, as the Gallic and Palmyrene Empires broke away from the Roman state, and a series of short-lived emperors led the empire. It was reunified under Aurelian. Diocletian set up the Tetrarchy, which divided the empire into four regions with two senior emperors (Augusti) and two junior emperors (Caesares). This successfully stabilized the empire, although it led to several civil wars in the early 4th century.

Constantine the Great emerged as sole emperor in 324. He founded Constantinople, which became the new capital of the empire, and was the first emperor to convert to Christianity. After the death of Theodosius I in 395, the empire was permanently divided between the East and the West."""

# -----------------
# 2. Write TXT
# -----------------
with open(os.path.join(DATA_DIR, "machine_learning.txt"), "w", encoding="utf-8") as f:
    f.write(txt_content)

# -----------------
# 3. Write PDF
# -----------------
def create_pdf(filename, text):
    doc = SimpleDocTemplate(filename, pagesize=letter)
    styles = getSampleStyleSheet()
    style = styles["Normal"]
    
    # SimpleDocTemplate needs a list of flowables
    story = []
    
    # Split text by paragraphs and add to story
    for para in text.split("\n\n"):
        p = Paragraph(para, style)
        story.append(p)
        story.append(Spacer(1, 12))
        
    doc.build(story)

create_pdf(os.path.join(DATA_DIR, "quantum_computing.pdf"), pdf_content)

# -----------------
# 4. Write DOCX
# -----------------
def create_docx(filename, text):
    doc = docx.Document()
    doc.add_heading('History of the Roman Empire', 0)
    
    for para in text.split("\n\n"):
        doc.add_paragraph(para)
        
    doc.save(filename)

create_docx(os.path.join(DATA_DIR, "roman_empire.docx"), docx_content)

print("Example documents generated successfully in data/ directory!")
